import os
import logging
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters
from telegram import Document
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from pdfminer.high_level import extract_text

# التوكن مال البوت (من BotFather)
TOKEN = os.environ.get("8367511871:AAFDNcsh3DlDPap4z2qKg6E6r1e0YUFHt7E")

# إعداد اللوج
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# دالة تحويل Word -> PDF
def word_to_pdf(input_file, output_file):
    doc = DocxDocument(input_file)
    c = canvas.Canvas(output_file, pagesize=letter)
    width, height = letter
    y = height - 50
    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            c.drawString(50, y, text)
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 50
    c.save()

# دالة تحويل PDF -> Word
def pdf_to_word(input_file, output_file):
    text = extract_text(input_file)
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_file)

# استقبال الملفات
def handle_file(update, context):
    file = update.message.document
    file_name = file.file_name.lower()

    if file_name.endswith(".docx"):
        file_path = "input.docx"
        output_file = "output.pdf"
        new_file = context.bot.getFile(file.file_id)
        new_file.download(file_path)
        word_to_pdf(file_path, output_file)
        update.message.reply_document(open(output_file, "rb"), filename=output_file)

    elif file_name.endswith(".pdf"):
        file_path = "input.pdf"
        output_file = "output.docx"
        new_file = context.bot.getFile(file.file_id)
        new_file.download(file_path)
        pdf_to_word(file_path, output_file)
        update.message.reply_document(open(output_file, "rb"), filename=output_file)

    else:
        update.message.reply_text("⚠️ أرسل ملف PDF أو Word فقط.")

def start(update, context):
    update.message.reply_text("👋 أهلاً! أرسل ملف PDF أو Word وسأحولّه لك.")

# تشغيل البوت
def main():
    updater = Updater(TOKEN, use_context=True)
    dp = updater.dispatcher
    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(MessageHandler(Filters.document, handle_file))

    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
